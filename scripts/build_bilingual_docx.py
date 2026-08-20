#!/usr/bin/env python3
"""Build a bilingual vocabulary-memory DOCX from audited sentence JSON."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


NAVY = RGBColor(31, 61, 99)
BLUE = RGBColor(35, 104, 168)
GREY = RGBColor(92, 101, 112)
LATIN_FONT = "Aptos"
EAST_ASIA_FONT = "Noto Sans CJK SC"


def set_east_asia_font(run, font_name: str = EAST_ASIA_FONT) -> None:
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), font_name)


def shade_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_bottom_border(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), "E8EDF3")
    borders.append(bottom)
    properties.append(borders)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def find_target_ranges(
    text: str,
    surfaces: list[str],
) -> tuple[list[tuple[int, int]], list[str]]:
    """Choose distinct, non-overlapping occurrences, preferring longer phrases."""
    unique = list(dict.fromkeys(surfaces))
    occupied: list[tuple[int, int]] = []
    unresolved: list[str] = []
    for surface in sorted(unique, key=lambda value: (-len(value), unique.index(value))):
        pattern = re.compile(
            rf"(?<![A-Za-z0-9]){re.escape(surface)}(?![A-Za-z0-9])",
            re.IGNORECASE,
        )
        candidates = [(match.start(), match.end()) for match in pattern.finditer(text)]
        chosen = next(
            (
                candidate
                for candidate in candidates
                if all(candidate[1] <= start or candidate[0] >= end for start, end in occupied)
            ),
            None,
        )
        if chosen is None:
            unresolved.append(surface)
        else:
            occupied.append(chosen)
    return sorted(occupied), unresolved


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(0)

    title = document.styles["Title"]
    title.font.name = LATIN_FONT
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = NAVY
    title._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)

    heading = document.styles["Heading 1"]
    heading.font.name = LATIN_FONT
    heading.font.size = Pt(14)
    heading.font.bold = True
    heading.font.color.rgb = NAVY
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(5)
    heading.paragraph_format.keep_with_next = True

    note = document.styles.add_style("Study Note", WD_STYLE_TYPE.PARAGRAPH)
    note.font.name = LATIN_FONT
    note.font.size = Pt(9.5)
    note.font.color.rgb = GREY
    note._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    note.paragraph_format.space_after = Pt(3)


def add_english_sentence(paragraph, sentence: dict[str, Any], display_index: int) -> tuple[int, list[str]]:
    number = paragraph.add_run(f"{display_index}. ")
    number.bold = True
    number.font.color.rgb = NAVY
    number.font.size = Pt(10.5)
    number.font.name = LATIN_FONT

    english = str(sentence["english"])
    surfaces = [str(target["surface"]) for target in sentence["targets"]]
    ranges, unresolved = find_target_ranges(english, surfaces)
    cursor = 0
    for start, end in ranges:
        if cursor < start:
            run = paragraph.add_run(english[cursor:start])
            run.font.name = LATIN_FONT
            run.font.size = Pt(11)
        run = paragraph.add_run(english[start:end])
        run.underline = True
        run.bold = True
        run.font.color.rgb = BLUE
        run.font.name = LATIN_FONT
        run.font.size = Pt(11)
        cursor = end
    if cursor < len(english):
        run = paragraph.add_run(english[cursor:])
        run.font.name = LATIN_FONT
        run.font.size = Pt(11)
    return len(ranges), unresolved


def build_document(
    data: list[dict[str, Any]],
    output_path: str | Path,
    title: str,
    version: str,
    scope_note: str,
) -> dict[str, Any]:
    if not data:
        raise ValueError("Sentence data must not be empty")
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    sentence_count = len(data)
    target_count = sum(len(item.get("targets", [])) for item in data)

    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(title)
    header_run.font.name = LATIN_FONT
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = GREY
    set_east_asia_font(header_run)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Page ")
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = GREY
    add_page_field(footer)

    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"{version} · Bilingual Study Edition · {target_count} targets · {sentence_count} sentences"
    )
    subtitle_run.font.name = LATIN_FONT
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = BLUE
    set_east_asia_font(subtitle_run)

    scope = document.add_paragraph(style="Study Note")
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.add_run("Scope: " + scope_note)
    shade_paragraph(scope, "F3F6F9")
    guide = document.add_paragraph(style="Study Note")
    guide.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide.add_run("Read the English, recall each underlined target, then check the Chinese.")
    guide.paragraph_format.space_after = Pt(8)

    current_scene = None
    underline_ranges = 0
    unresolved_all: list[str] = []
    for display_index, item in enumerate(data, start=1):
        scene = str(item.get("scene") or "Vocabulary Sentences")
        if scene != current_scene:
            current_scene = scene
            heading = document.add_paragraph(scene, style="Heading 1")
            shade_paragraph(heading, "EEF3F8")

        english_paragraph = document.add_paragraph()
        english_paragraph.paragraph_format.keep_with_next = True
        english_paragraph.paragraph_format.space_before = Pt(2)
        english_paragraph.paragraph_format.space_after = Pt(1)
        english_paragraph.paragraph_format.line_spacing = 1.08
        count, unresolved = add_english_sentence(english_paragraph, item, display_index)
        underline_ranges += count
        unresolved_all.extend(
            f'{item.get("sentence_id", "?")}:{surface}' for surface in unresolved
        )

        chinese_paragraph = document.add_paragraph()
        chinese_paragraph.paragraph_format.space_after = Pt(5)
        chinese_paragraph.paragraph_format.keep_together = True
        chinese_paragraph.paragraph_format.widow_control = True
        chinese_run = chinese_paragraph.add_run("中文：" + str(item["chinese"]))
        chinese_run.font.name = LATIN_FONT
        chinese_run.font.size = Pt(10)
        chinese_run.font.color.rgb = GREY
        set_east_asia_font(chinese_run)
        add_bottom_border(chinese_paragraph)

    if unresolved_all:
        raise ValueError("Unresolved target surfaces: " + ", ".join(unresolved_all))

    core = document.core_properties
    core.title = f"{title} {version}"
    core.subject = "Audited bilingual vocabulary memory sentences"
    core.author = "vocabulary-memory-sentences-skill"
    core.keywords = "vocabulary, bilingual, memory sentences, study"
    document.save(output)
    return {
        "output": str(output),
        "sentences": sentence_count,
        "targets": target_count,
        "underline_ranges": underline_ranges,
        "unresolved": len(unresolved_all),
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--data", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--title", default="Vocabulary Memory Sentences")
    parser.add_argument("--version", default="v0.1.0")
    parser.add_argument("--scope-note", default="Visible source entries only.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    data = json.loads(args.data.read_text(encoding="utf-8"))
    result = build_document(data, args.output, args.title, args.version, args.scope_note)
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
